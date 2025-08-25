import os
import io
import json
import uuid
import hashlib
import mimetypes
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
from docx import Document
import openai
from openai import OpenAI
import whisper

class FileProcessingService:
    def __init__(self):
        self.client = OpenAI()
        self.whisper_model = None
        self.supported_document_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'text/csv': 'csv'
        }
        self.supported_audio_types = {
            'audio/wav': 'wav',
            'audio/mpeg': 'mp3',
            'audio/mp4': 'm4a',
            'audio/x-m4a': 'm4a',
            'audio/webm': 'webm',
            'audio/ogg': 'ogg'
        }
        self.max_file_size = 100 * 1024 * 1024  # 100MB limit
        
        # Initialize Whisper model for audio transcription
        self._initialize_whisper()
    
    def _initialize_whisper(self):
        """Initialize Whisper model for audio transcription"""
        try:
            # Use smaller model for faster processing
            self.whisper_model = whisper.load_model("base")
        except Exception as e:
            print(f"Failed to initialize Whisper model: {e}")
            self.whisper_model = None
    
    def process_file(self, file, user_id: str) -> Dict[str, Any]:
        """Process uploaded file with AI analysis"""
        try:
            # Validate file
            is_valid, validation_message = self.validate_file(file)
            if not is_valid:
                raise Exception(validation_message)
            
            # Generate file ID and metadata
            file_id = str(uuid.uuid4())
            file_hash = self._calculate_file_hash(file)
            
            # Extract text content
            text_content = self.extract_text_from_file(file)
            
            # Analyze content with AI
            analysis = self.analyze_content_with_ai(text_content, file)
            
            # Process for RAG (if enabled)
            rag_document_id = None
            if analysis.get('extracted_text'):
                rag_document_id = self._process_for_rag(
                    analysis['extracted_text'], 
                    file.filename, 
                    user_id
                )
            
            # Create result
            result = {
                'file_id': file_id,
                'filename': file.filename,
                'file_size': len(file.read()),
                'file_hash': file_hash,
                'mime_type': file.content_type,
                'user_id': user_id,
                'upload_timestamp': datetime.utcnow().isoformat(),
                'processing_status': 'completed',
                'analysis': analysis,
                'rag_document_id': rag_document_id
            }
            
            # Reset file pointer for potential future use
            file.seek(0)
            
            return result
            
        except Exception as e:
            raise Exception(f"File processing failed: {str(e)}")
    
    def validate_file(self, file) -> Tuple[bool, str]:
        """Validate uploaded file for security and format compliance"""
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > self.max_file_size:
            return False, f"File size exceeds maximum limit of {self.max_file_size // (1024*1024)}MB"
        
        # Check if file type is supported
        content_type = file.content_type
        if content_type not in self.supported_document_types and content_type not in self.supported_audio_types:
            return False, f"Unsupported file type: {content_type}"
        
        # Basic security checks
        dangerous_extensions = ['.exe', '.bat', '.cmd', '.scr', '.pif', '.com', '.vbs', '.js']
        file_ext = os.path.splitext(file.filename.lower())[1]
        if file_ext in dangerous_extensions:
            return False, "Potentially dangerous file type detected"
        
        return True, "File validation passed"
    
    def extract_text_from_file(self, file) -> str:
        """Extract text content from various file types"""
        try:
            content_type = file.content_type
            
            if content_type in self.supported_document_types:
                return self._extract_from_document(file, content_type)
            elif content_type in self.supported_audio_types:
                return self._extract_from_audio(file)
            else:
                raise Exception(f"Unsupported file type: {content_type}")
                
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _extract_from_document(self, file, content_type: str) -> str:
        """Extract text from document files"""
        try:
            if content_type == 'application/pdf':
                return self._extract_from_pdf(file)
            elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._extract_from_docx(file)
            elif content_type in ['text/plain', 'text/markdown', 'text/csv']:
                return self._extract_from_text(file)
            else:
                raise Exception(f"Unsupported document type: {content_type}")
                
        except Exception as e:
            raise Exception(f"Document extraction failed: {str(e)}")
    
    def _extract_from_pdf(self, file) -> str:
        """Extract text content from PDF file"""
        try:
            pdf_file = io.BytesIO(file.read())
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file) -> str:
        """Extract text content from DOCX file"""
        try:
            docx_file = io.BytesIO(file.read())
            doc = Document(docx_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_from_text(self, file) -> str:
        """Extract text content from plain text files"""
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    return file.read().decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Unable to decode text file with supported encodings")
            
        except Exception as e:
            raise Exception(f"Failed to extract text from file: {str(e)}")
    
    def _extract_from_audio(self, file) -> str:
        """Extract text from audio files using Whisper"""
        try:
            if not self.whisper_model:
                raise Exception("Whisper model not initialized")
            
            # Save audio to temporary file
            temp_path = f"/tmp/audio_{uuid.uuid4()}.wav"
            file.save(temp_path)
            
            try:
                # Transcribe audio
                result = self.whisper_model.transcribe(temp_path)
                transcription = result["text"]
                
                return transcription.strip()
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        except Exception as e:
            raise Exception(f"Audio transcription failed: {str(e)}")
    
    def analyze_content_with_ai(self, text_content: str, file) -> Dict[str, Any]:
        """Analyze content using OpenAI GPT models"""
        try:
            # Prepare content for analysis (limit to avoid token limits)
            analysis_text = text_content[:4000] if len(text_content) > 4000 else text_content
            
            # Create analysis prompt
            prompt = f"""Analyze the following document content and provide comprehensive insights:

Document: {file.filename}
Type: {file.content_type}
Content Length: {len(text_content)} characters
Content Preview: {analysis_text[:1000]}...

Please provide:
1. A concise summary (2-3 sentences)
2. Key points and insights (5-7 bullet points)
3. Document category and tags
4. Estimated reading time
5. Content quality assessment
6. Suggested actions or follow-ups

Format as JSON with keys: summary, key_points, category, tags, reading_time_minutes, quality_score, suggested_actions"""

            # Get AI analysis
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert document analyst. Provide insightful analysis of documents."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.3
            )
            
            ai_response = response.choices[0].message.content
            
            # Try to parse AI response
            try:
                analysis_data = json.loads(ai_response)
            except:
                # If parsing fails, create fallback analysis
                analysis_data = self._create_fallback_analysis(text_content, file)
            
            # Add extracted text and metadata
            analysis_data['extracted_text'] = text_content
            analysis_data['document_type'] = self._get_document_type(file.content_type)
            analysis_data['word_count'] = len(text_content.split())
            analysis_data['sentence_count'] = len(text_content.split('.'))
            analysis_data['paragraph_count'] = len(text_content.split('\n\n'))
            analysis_data['content_length'] = len(text_content)
            analysis_data['ai_model_used'] = 'gpt-3.5-turbo'
            analysis_data['analysis_timestamp'] = datetime.utcnow().isoformat()
            
            return analysis_data
            
        except Exception as e:
            print(f"AI analysis failed: {e}")
            # Return fallback analysis
            return self._create_fallback_analysis(text_content, file)
    
    def _create_fallback_analysis(self, text_content: str, file) -> Dict[str, Any]:
        """Create fallback analysis when AI fails"""
        word_count = len(text_content.split())
        reading_time = max(1, word_count // 200)  # Assume 200 words per minute
        
        return {
            'summary': f"Document '{file.filename}' contains {word_count} words of content.",
            'key_points': [
                f"Document type: {self._get_document_type(file.content_type)}",
                f"Content length: {len(text_content)} characters",
                f"Estimated reading time: {reading_time} minutes"
            ],
            'category': 'general',
            'tags': [self._get_document_type(file.content_type), 'analyzed'],
            'reading_time_minutes': reading_time,
            'quality_score': 0.7,
            'suggested_actions': ['Review content', 'Extract key information', 'Store for future reference']
        }
    
    def _get_document_type(self, content_type: str) -> str:
        """Get human-readable document type"""
        if content_type in self.supported_document_types:
            return self.supported_document_types[content_type].upper()
        elif content_type in self.supported_audio_types:
            return 'AUDIO'
        else:
            return 'UNKNOWN'
    
    def _calculate_file_hash(self, file) -> str:
        """Calculate SHA-256 hash of file content"""
        try:
            file.seek(0)
            file_content = file.read()
            file_hash = hashlib.sha256(file_content).hexdigest()
            file.seek(0)  # Reset file pointer
            return file_hash
        except Exception as e:
            print(f"Failed to calculate file hash: {e}")
            return "unknown"
    
    def _process_for_rag(self, text_content: str, filename: str, user_id: str) -> Optional[str]:
        """Process document for RAG (Retrieval-Augmented Generation)"""
        try:
            # This would integrate with the RAG service
            # For now, return a placeholder
            return f"rag_doc_{uuid.uuid4()}"
        except Exception as e:
            print(f"RAG processing failed: {e}")
            return None
    
    def transcribe_audio(self, audio_file, user_id: str) -> str:
        """Transcribe audio file using Whisper"""
        try:
            if not self.whisper_model:
                raise Exception("Whisper model not initialized")
            
            # Save audio to temporary file
            temp_path = f"/tmp/audio_{uuid.uuid4()}.wav"
            audio_file.save(temp_path)
            
            try:
                # Transcribe audio
                result = self.whisper_model.transcribe(temp_path)
                transcription = result["text"]
                
                return transcription.strip()
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        except Exception as e:
            raise Exception(f"Audio transcription failed: {str(e)}")
    
    def get_processing_stats(self, user_id: str) -> Dict[str, Any]:
        """Get processing statistics for a user"""
        try:
            # This would query the database for actual stats
            # For now, return mock data
            return {
                'total_files_processed': 0,
                'files_by_type': {},
                'total_processing_time': 0,
                'success_rate': 1.0,
                'last_processed': None
            }
        except Exception as e:
            print(f"Failed to get processing stats: {e}")
            return {}
